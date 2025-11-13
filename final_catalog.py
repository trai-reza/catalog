#!/usr/bin/env python3
"""Generate final PDF catalog with better image-content matching"""
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib import colors
from PIL import Image as PILImage
import os
import tempfile
import shutil

def extract_images_with_descriptions(docx_path):
    """Extract images and try to find their descriptions"""
    doc = Document(docx_path)
    images = []
    temp_dir = tempfile.mkdtemp()
    
    # Extract images
    image_paths = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_ext = os.path.splitext(rel.target_ref)[1] or '.png'
                image_path = os.path.join(temp_dir, f"image_{i}{image_ext}")
                with open(image_path, 'wb') as f:
                    f.write(image_part.blob)
                image_paths.append((i, image_path))
            except Exception as e:
                print(f"Error extracting image {i}: {e}")
    
    # Try to extract descriptions - images in docx are often in separate paragraphs
    # Look for text paragraphs that might describe images
    descriptions = []
    para_texts = [p.text.strip() for p in doc.paragraphs]
    
    # Simple heuristic: descriptions are often short paragraphs near image locations
    # or paragraphs that are not headings
    for i, (orig_idx, img_path) in enumerate(image_paths):
        desc = ""
        # Look in paragraphs for potential descriptions
        # In Word, images and their captions are often together
        if i < len(para_texts):
            # Check a few paragraphs around where image might be
            for j in range(max(0, i*5), min(len(para_texts), (i+1)*5 + 3)):
                text = para_texts[j]
                if text and 10 < len(text) < 200 and not text.isupper():
                    desc = text
                    break
        descriptions.append(desc)
    
    result = [(img_path, desc) for (_, img_path), desc in zip(image_paths, descriptions)]
    return result, temp_dir

def get_image_dimensions(image_path, max_width=5.5*inch, max_height=7*inch):
    """Get image dimensions scaled to fit"""
    try:
        img = PILImage.open(image_path)
        width, height = img.size
        aspect = width / height
        
        if width > max_width:
            width = max_width
            height = width / aspect
        
        if height > max_height:
            height = max_height
            width = height * aspect
        
        return width, height
    except:
        return 4*inch, 3*inch

def create_catalog():
    print("Creating final catalog PDF...")
    
    # Extract content
    print("Reading main-content.docx...")
    main_doc = Document('main-content.docx')
    main_content = [para.text for para in main_doc.paragraphs if para.text.strip()]
    print(f"Found {len(main_content)} content paragraphs")
    
    print("Extracting images from photos.docx...")
    images_with_desc, temp_dir = extract_images_with_descriptions('photos.docx')
    print(f"Found {len(images_with_desc)} images")
    
    # Separate biography image (last one)
    biography_image = images_with_desc[-1] if images_with_desc else (None, "")
    regular_images = images_with_desc[:-1] if len(images_with_desc) > 1 else []
    
    # Create PDF
    pdf_path = 'catalog.pdf'
    doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                           rightMargin=0.7*inch, leftMargin=0.7*inch,
                           topMargin=0.7*inch, bottomMargin=0.7*inch)
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Graphical title style
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=26,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=15,
        spaceBefore=10,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Section heading with background
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=15,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=10,
        spaceBefore=14,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=5,
        alignment=TA_JUSTIFY,
        leading=12
    )
    
    image_desc_style = ParagraphStyle(
        'ImageDesc',
        parent=styles['BodyText'],
        fontSize=9,
        textColor=colors.HexColor('#7f8c8d'),
        spaceAfter=8,
        alignment=TA_CENTER,
        fontStyle='Italic',
        leading=10
    )
    
    # Build story
    story = []
    
    # Title page
    story.append(Spacer(1, 1*inch))
    
    # Logo
    if os.path.exists('LOGO.jpg'):
        try:
            logo = Image('LOGO.jpg', width=2.5*inch, height=2.5*inch)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 0.4*inch))
        except:
            pass
    
    # Main title with graphical styling
    if main_content:
        title_text = main_content[0] if len(main_content[0]) < 120 else "Film Catalog"
        title_html = f'<para backColor="#ecf0f1" borderWidth="3" borderColor="#2c3e50" borderPadding="15"><font size="26" color="#2c3e50"><b>{title_text}</b></font></para>'
        story.append(Paragraph(title_html, title_style))
        
        # Subtitle
        if len(main_content) > 1:
            subtitle_text = main_content[1]
            if len(subtitle_text) < 150:
                subtitle_html = f'<font size="13" color="#7f8c8d"><i>{subtitle_text}</i></font>'
                subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], 
                                               alignment=TA_CENTER, spaceAfter=20)
                story.append(Paragraph(subtitle_html, subtitle_style))
    
    story.append(PageBreak())
    
    # Process all content with images distributed throughout
    image_index = 0
    content_index = 0
    
    # Calculate image distribution - spread images evenly
    total_content = len(main_content)
    if total_content > 0 and len(regular_images) > 0:
        images_per_section = max(1, len(regular_images) // max(1, (total_content // 15)))
    else:
        images_per_section = 1
    
    section_keywords = ['title', 'director', 'writer', 'producer', 'production', 'cinematographer',
                       'synopsis', 'story', 'plot', 'cast', 'crew', 'credits', 'festival', 
                       'award', 'review', 'biography', 'bio']
    
    while content_index < len(main_content):
        text = main_content[content_index]
        text_lower = text.lower()
        
        # Determine if this is a heading
        is_heading = False
        if content_index < 10:  # Early paragraphs are often headings
            is_heading = True
        elif any(keyword in text_lower for keyword in section_keywords):
            is_heading = True
        elif len(text) < 70 and (text.isupper() or ':' in text or text.endswith(':')):
            is_heading = True
        
        # Add heading with graphical style
        if is_heading:
            heading_html = f'<para backColor="#ecf0f1" borderPadding="10" leftIndent="5"><font size="15" color="#2c3e50"><b>{text}</b></font></para>'
            story.append(Paragraph(heading_html, heading_style))
        else:
            # Regular content
            story.append(Paragraph(text, body_style))
        
        content_index += 1
        
        # Add images at strategic points
        if image_index < len(regular_images):
            # Add image after certain sections or every N paragraphs
            should_add = False
            
            if any(keyword in text_lower for keyword in ['synopsis', 'story', 'plot', 'description']):
                should_add = True
            elif content_index % 12 == 0 and content_index > 10:  # Every 12 paragraphs
                should_add = True
            elif content_index == len(main_content) - 5:  # Near end (before biography)
                should_add = True
            
            if should_add:
                img_path, img_desc = regular_images[image_index]
                try:
                    img_width, img_height = get_image_dimensions(img_path)
                    img = Image(img_path, width=img_width, height=img_height)
                    img.hAlign = 'CENTER'
                    story.append(Spacer(1, 0.25*inch))
                    story.append(img)
                    
                    # Add description if available
                    if img_desc and img_desc.strip():
                        story.append(Spacer(1, 0.1*inch))
                        desc_html = f'<i>{img_desc}</i>'
                        story.append(Paragraph(desc_html, image_desc_style))
                    
                    story.append(Spacer(1, 0.25*inch))
                    image_index += 1
                except Exception as e:
                    print(f"Error adding image {image_index}: {e}")
                    image_index += 1
        
        # Small spacing
        if content_index % 5 == 0:
            story.append(Spacer(1, 0.05*inch))
    
    # Add any remaining images
    while image_index < len(regular_images):
        img_path, img_desc = regular_images[image_index]
        try:
            story.append(Spacer(1, 0.3*inch))
            img_width, img_height = get_image_dimensions(img_path)
            img = Image(img_path, width=img_width, height=img_height)
            img.hAlign = 'CENTER'
            story.append(img)
            if img_desc and img_desc.strip():
                story.append(Spacer(1, 0.1*inch))
                desc_html = f'<i>{img_desc}</i>'
                story.append(Paragraph(desc_html, image_desc_style))
            story.append(Spacer(1, 0.25*inch))
            image_index += 1
        except Exception as e:
            print(f"Error adding remaining image: {e}")
            image_index += 1
    
    # Biography section
    story.append(PageBreak())
    
    # Biography heading with prominent style
    bio_heading_html = '<para backColor="#2c3e50" borderPadding="14"><font size="18" color="white"><b>Biography</b></font></para>'
    bio_heading_style = ParagraphStyle('BioHeading', parent=styles['Heading1'], 
                                       alignment=TA_CENTER, spaceAfter=20, spaceBefore=10)
    story.append(Paragraph(bio_heading_html, bio_heading_style))
    
    # Biography image
    if biography_image[0]:
        try:
            img_path = biography_image[0]
            img_width, img_height = get_image_dimensions(img_path, max_width=5*inch, max_height=6*inch)
            img = Image(img_path, width=img_width, height=img_height)
            img.hAlign = 'CENTER'
            story.append(Spacer(1, 0.2*inch))
            story.append(img)
            
            # Biography image description
            if biography_image[1] and biography_image[1].strip():
                story.append(Spacer(1, 0.15*inch))
                desc_html = f'<i>{biography_image[1]}</i>'
                story.append(Paragraph(desc_html, image_desc_style))
            
            story.append(Spacer(1, 0.3*inch))
        except Exception as e:
            print(f"Error adding biography image: {e}")
    
    # Look for biography-related text in main content
    for i, text in enumerate(main_content):
        if 'biography' in text.lower() or 'bio' in text.lower():
            # Add following paragraphs that might be biography content
            for j in range(i + 1, min(i + 15, len(main_content))):
                bio_text = main_content[j]
                # Stop if we hit another heading
                if any(kw in bio_text.lower() for kw in section_keywords) and len(bio_text) < 80:
                    break
                story.append(Paragraph(bio_text, body_style))
            break
    
    # Build PDF
    print("Building PDF...")
    doc.build(story)
    print(f"✓ Catalog saved to {pdf_path}")
    print(f"  - {len(main_content)} content paragraphs included")
    print(f"  - {len(images_with_desc)} images included")
    
    # Cleanup
    shutil.rmtree(temp_dir, ignore_errors=True)
    
    return pdf_path

if __name__ == '__main__':
    create_catalog()
